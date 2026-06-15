from docx import Document
import json


class ReportGenerator:

    def generate_report(
        self,
        doc_result,
        impact_result,
        output_file="comparison_report.docx"
    ):
        doc = Document()

        # ✅ Title
        doc.add_heading("AI Document Comparison Report", level=1)

        # ✅ Document Analysis
        doc.add_heading("Document Analysis", level=2)
        print("Hiiii")
        print(type(doc_result))


        if isinstance(doc_result, str):
            try:
                doc_result = json.loads(doc_result)
            except Exception:
                doc_result = []

        if isinstance(doc_result, dict):
            doc_result = [doc_result]
        print(type(doc_result))

        for idx, item in enumerate(doc_result, start=1):
            # If item is a string, try to parse it
            if isinstance(item, str):
                try:
                    item = json.loads(item)
                except Exception:
                    item = {}

            doc.add_heading(f"Rule {idx}", level=3)
            doc.add_paragraph(f"Change Type: {item.get('change_type', 'N/A')}")
            doc.add_paragraph(f"Summary: {item.get('summary', 'N/A')}")
            doc.add_paragraph(f"Rule Name: {item.get('rule_name', 'N/A')}")

            # ✅ Parse "changes" if it's a JSON string
            changes = item.get("changes", {})
            if isinstance(changes, str):
                try:
                    changes = json.loads(changes)
                except Exception:
                    changes = {}

            # Handle dict of numbered rules OR old/new chunks
            if isinstance(changes, dict):
                if "old_chunk" in changes or "new_chunk" in changes:
                    doc.add_paragraph(f"Old Text: {changes.get('old_chunk', 'N/A')}")
                    doc.add_paragraph(f"New Text: {changes.get('new_chunk', 'N/A')}")
                else:
                    # ✅ Iterate over numbered rules and add them as bullet points
                    for key, value in changes.items():
                        doc.add_paragraph(value, style="List Bullet")
            else:
                doc.add_paragraph("No detailed changes available")

        # ✅ Impact Analysis
        doc.add_page_break()
        doc.add_heading("Impact Analysis", level=2)

        if isinstance(impact_result, str):
            try:
                impact_result = json.loads(impact_result)
            except Exception:
                impact_result = {}

        fields = [
            "what_changed",
            "business_impact",
            "compliance_impact",
            "stakeholders_affected",
            "risk_level",
            "recommended_actions",
            "executive_summary"
        ]

        for field in fields:
            doc.add_heading(field.replace("_", " ").title(), level=3)
            doc.add_paragraph(str(impact_result.get(field, "N/A")))

        doc.save(output_file)
        print("Report generated successfully:", output_file)

        return output_file
